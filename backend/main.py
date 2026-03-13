from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import shutil
import os
from pathlib import Path
from docling.document_converter import DocumentConverter
from docx import Document
from docx.shared import Pt, Inches
import pandas as pd
import camelot
from paddleocr import PaddleOCR
import io
import fitz  # PyMuPDF for scan detection

app = FastAPI(title="Nexora Advanced PDF Engine")

# Enable CORS for the Vue frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize OCR
ocr = PaddleOCR(use_angle_cls=True, lang='es')

def is_pdf_scanned(pdf_path):
    """Detects if a PDF is scanned by checking for actual text."""
    doc = fitz.open(pdf_path)
    for page in doc:
        if page.get_text().strip():
            return False
    return True

@app.post("/convert/word")
async def convert_to_word(file: UploadFile = File(...)):
    temp_path = Path(f"temp_{file.filename}")
    try:
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # 1. Check if scanned
        is_scanned = is_pdf_scanned(str(temp_path))
        
        # 2. Extract content
        doc = Document()
        
        if is_scanned:
            # Apply PaddleOCR
            result = ocr.ocr(str(temp_path), cls=True)
            for page in result:
                for line in page:
                    text = line[1][0]
                    doc.add_paragraph(text)
        else:
            # Structural Analysis with Docling
            converter = DocumentConverter()
            conversion_result = converter.convert(str(temp_path))
            doc_content = conversion_result.document
            
            # Reconstruct elements
            for element in doc_content.export_to_dict().get("elements", []):
                if element["type"] == "heading":
                    doc.add_heading(element["text"], level=element.get("level", 1))
                elif element["type"] == "paragraph":
                    p = doc.add_paragraph()
                    run = p.add_run(element["text"])
                    if element.get("bold"): run.bold = True
                elif element["type"] == "table":
                    # For Tables, we could also use Camelot here for higher precision
                    pass

        # 3. Handle Tables with Camelot (Advanced)
        tables = camelot.read_pdf(str(temp_path), pages='all', flavor='stream')
        for table in tables:
            df = table.df
            t = doc.add_table(rows=df.shape[0], cols=df.shape[1])
            t.style = 'Table Grid'
            for i, row in enumerate(df.values):
                for j, val in enumerate(row):
                    t.cell(i, j).text = str(val)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return StreamingResponse(
            output, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            headers={"Content-Disposition": f"attachment; filename={file.filename.replace('.pdf', '')}.docx"}
        )

    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if temp_path.exists():
            os.remove(temp_path)

@app.post("/convert/excel")
async def convert_to_excel(file: UploadFile = File(...)):
    temp_path = Path(f"temp_ex_{file.filename}")
    try:
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        tables = camelot.read_pdf(str(temp_path), pages='all')
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for i, table in enumerate(tables):
                table.df.to_excel(writer, sheet_name=f'Tabla {i+1}', index=False)
        
        output.seek(0)
        return StreamingResponse(
            output, 
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
            headers={"Content-Disposition": f"attachment; filename={file.filename.replace('.pdf', '')}.xlsx"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if temp_path.exists():
            os.remove(temp_path)

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
